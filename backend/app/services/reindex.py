from __future__ import annotations

import logging
import os
from typing import Any, BinaryIO, List

from docx import Document

from app.db.chroma import get_collection
from app.services.rag import upsert_proverbs


logger = logging.getLogger(__name__)


def read_docx_lines(path: str) -> List[str]:
    """Read non-empty, stripped paragraphs from a .docx file."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    doc = Document(path)
    lines: List[str] = []
    for p in doc.paragraphs:
        if p is None:
            continue
        text = (p.text or "").strip()
        if not text:
            continue
        lines.append(text)
    return lines


def read_docx_lines_from_stream(file_stream: BinaryIO) -> List[str]:
    """Read non-empty, stripped paragraphs from a .docx file stream."""
    file_stream.seek(0)
    doc = Document(file_stream)
    lines: List[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        lines.append(text)
    return lines


def merge_proverbs_meanings(proverbs_path: str, meanings_path: str) -> list[dict[str, Any]]:
    """Merge two .docx files (proverbs & meanings) into rows suitable for upsert.

    Each row: {keyword: "", proverb: str, meaning: str, example: ""}
    """
    proverbs = read_docx_lines(proverbs_path)
    meanings = read_docx_lines(meanings_path)

    rows: list[dict[str, Any]] = []

    if len(proverbs) != len(meanings):
        logger.warning(
            "Mismatch in lengths: proverbs=%d meanings=%d — extra items will be ignored",
            len(proverbs),
            len(meanings),
        )

    for p, m in zip(proverbs, meanings):
        rows.append({"keyword": "", "proverb": p, "meaning": m, "example": ""})

    return rows


def merge_proverbs_meanings_from_uploads(
    proverbs_file: Any, meanings_file: Any
) -> list[dict[str, Any]]:
    """Merge two uploaded .docx files into rows suitable for upsert."""
    proverbs = read_docx_lines_from_stream(proverbs_file.file)
    meanings = read_docx_lines_from_stream(meanings_file.file)

    rows: list[dict[str, Any]] = []

    if len(proverbs) != len(meanings):
        logger.warning(
            "Mismatch in uploaded lengths: proverbs=%d meanings=%d — extra items will be ignored",
            len(proverbs),
            len(meanings),
        )

    for p, m in zip(proverbs, meanings):
        rows.append({"keyword": "", "proverb": p, "meaning": m, "example": ""})

    return rows


def _clear_collection(col: Any) -> None:
    """Remove all records from the Chroma collection by deleting all IDs."""
    result = col.get(limit=100000, include=["metadatas"])
    ids = result.get("ids") or []
    if ids:
        col.delete(ids=ids)


def reindex_from_word_files(proverbs_path: str, meanings_path: str, clear_existing: bool = True) -> dict[str, Any]:
    """Rebuild the ChromaDB collection from two Word files.

    Returns dict with inserted and skipped counts and warnings.
    """
    col = get_collection()

    rows = merge_proverbs_meanings(proverbs_path, meanings_path)

    warnings: list[str] = []

    if not rows:
        warnings.append("No valid rows found after merge.")
        return {"inserted": 0, "skipped": 0, "warnings": warnings}

    if clear_existing:
        try:
            _clear_collection(col)
        except Exception as e:
            logger.exception("Failed to clear existing ChromaDB collection: %s", e)
            warnings.append(f"Failed to clear existing collection: {e}")

    inserted, skipped = upsert_proverbs(rows)

    return {"inserted": inserted, "skipped": skipped, "warnings": warnings}


def reindex_from_word_uploads(
    proverbs_file: Any, meanings_file: Any, clear_existing: bool = True
) -> dict[str, Any]:
    """Rebuild the ChromaDB collection from uploaded .docx files."""
    col = get_collection()

    rows = merge_proverbs_meanings_from_uploads(proverbs_file, meanings_file)

    warnings: list[str] = []

    if not rows:
        warnings.append("No valid rows found after merge.")
        return {"inserted": 0, "skipped": 0, "warnings": warnings}

    if clear_existing:
        try:
            _clear_collection(col)
        except Exception as e:
            logger.exception("Failed to clear existing ChromaDB collection: %s", e)
            warnings.append(f"Failed to clear existing collection: {e}")

    inserted, skipped = upsert_proverbs(rows)

    return {"inserted": inserted, "skipped": skipped, "warnings": warnings}
