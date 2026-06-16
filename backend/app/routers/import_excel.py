import io

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile
from openpyxl import load_workbook
from docx import Document

from app.core.config import settings
from app.core.deps import get_current_user_id, require_admin
from app.services.rag import upsert_proverbs


router = APIRouter()

REQUIRED_COLUMNS = {"keyword", "proverb", "meaning", "example"}


def _read_excel_rows(raw: bytes) -> list[dict[str, str]]:
    """Read .xlsx into list of row dicts using openpyxl (no pandas needed)."""
    wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    ws = wb.active

    rows_iter = ws.iter_rows(values_only=True)
    try:
        header_row = next(rows_iter)
    except StopIteration:
        wb.close()
        return []

    headers = [str(h).strip() if h is not None else "" for h in header_row]
    missing = REQUIRED_COLUMNS - set(headers)
    if missing:
        wb.close()
        raise ValueError(f"Missing columns: {sorted(missing)}")

    records: list[dict[str, str]] = []
    for row in rows_iter:
        if not any(cell is not None and str(cell).strip() for cell in row):
            continue
        record: dict[str, str] = {}
        for i, header in enumerate(headers):
            if not header:
                continue
            value = row[i] if i < len(row) else None
            record[header] = str(value).strip() if value is not None else ""
        records.append(record)

    wb.close()
    return records


def _read_docx_rows(raw: bytes) -> list[dict[str, str]]:
    """
    Read .docx where the first table contains a header row with:
    keyword | proverb | meaning | example
    """
    doc = Document(io.BytesIO(raw))
    if not doc.tables:
        return []

    table = doc.tables[0]
    if not table.rows:
        return []

    header_cells = [c.text.strip() for c in table.rows[0].cells]
    missing = REQUIRED_COLUMNS - set(header_cells)
    if missing:
        raise ValueError(f"Missing columns: {sorted(missing)}")

    records: list[dict[str, str]] = []
    for r in table.rows[1:]:
        cells = [c.text.strip() for c in r.cells]
        if not any(cells):
            continue
        record: dict[str, str] = {}
        for i, header in enumerate(header_cells):
            if not header:
                continue
            record[header] = cells[i] if i < len(cells) else ""
        records.append(record)
    return records


@router.post("/import-excel")
async def import_excel(
    file: UploadFile = File(...),
    _admin=Depends(require_admin),
):
    if not file.filename or not file.filename.lower().endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Please upload a .xlsx file")

    raw = await file.read()
    try:
        rows = _read_excel_rows(raw)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read Excel: {e}")

    inserted, skipped = upsert_proverbs(rows)

    return {"inserted": inserted, "skipped": skipped, "collection": settings.chroma_collection_name}


@router.post("/import-docx")
async def import_docx(
    file: UploadFile = File(...),
    _admin=Depends(require_admin),
):
    if not file.filename or not file.filename.lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Please upload a .docx file")

    raw = await file.read()
    try:
        rows = _read_docx_rows(raw)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read DOCX: {e}")

    inserted, skipped = upsert_proverbs(rows)
    return {"inserted": inserted, "skipped": skipped, "collection": settings.chroma_collection_name}
